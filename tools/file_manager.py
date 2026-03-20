"""文件管理工具：用于各种文件的增删查改，支持多种文件格式，在沙盒中运行。"""

from __future__ import annotations

import json
import os
import base64
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_core.tools import tool

from utils.path import ensure_task_dirs, get_task_dir
from utils.task_context import get_task_id

# 导入文件处理库
from docx import Document
import json as json_module



def _ensure_sandbox_path(task_id: str, file_path: str) -> Path:
    """
    确保文件路径在沙盒中，返回绝对路径
    
    Args:
        task_id: 任务ID
        file_path: 文件路径（相对或绝对）
    
    Returns:
        沙盒中的绝对文件路径
    """
    task_dir = ensure_task_dirs(task_id)
    
    # 如果是绝对路径，提取文件名和扩展名，在沙盒中创建
    if Path(file_path).is_absolute():
        file_name = Path(file_path).name
        return task_dir / file_name
    # 如果是相对路径，直接在沙盒中创建
    else:
        return task_dir / file_path


def _validate_file_extension(file_path: str) -> bool:
    """
    验证文件扩展名是否支持
    
    Args:
        file_path: 文件路径
    
    Returns:
        是否支持该文件格式
    """
    supported_extensions = [
        '.doc', '.docx', '.txt', '.json', '.md',
        '.csv', '.py', '.html', '.htm', '.css', '.js', '.pdf', '.xlsx'
    ]
    ext = Path(file_path).suffix.lower()
    return ext in supported_extensions


def _sandbox_relative_path(task_id: str, sandbox_path: Path) -> str:
    task_dir = get_task_dir(task_id).resolve()
    try:
        rel = sandbox_path.resolve().relative_to(task_dir)
    except Exception:
        return ""
    rel_str = str(rel).replace("\\", "/")
    if not rel_str:
        return f"sandbox/{task_id}"
    return f"sandbox/{task_id}/{rel_str}"


@tool
def file_manager(
    action: str,
    file_path: str,
    content: Optional[str] = None,
    data: Optional[Dict] = None
) -> str:
    """
    工具名称：file_manager
    描述：文件管理工具，用于在沙盒中执行文件的增删查改操作，支持多种文件格式（doc、docx、txt、json、md、csv）。
    强制使用：当用户提出任何与文件相关的请求时，**绝对必须**使用本工具，**不允许**直接模拟文件操作。
    使用时机：以下情况**必须**调用本工具：
    - 用户要求查看、读取、检查任何文件内容
    - 用户要求创建新文件或写入内容
    - 用户要求修改、更新、编辑现有文件
    - 用户要求删除文件
    - 用户要求列出目录中的文件
    - 用户要求确认文件是否存在
    输入参数：
    - action（必填）：操作类型，**必须**是以下值之一：
        - "read"：读取文件内容
        - "create"：创建新文件
        - "update"：更新现有文件
        - "delete"：删除文件
        - "list"：列出目录内容
    - file_path（必填）：文件或目录路径（相对或绝对路径，会自动转换为沙盒路径）
    - content（必填）：当action为"create"或"update"且文件非JSON格式时，**必须**提供文件内容
    - data（必填）：当action为"create"或"update"且文件为JSON格式时，**必须**提供JSON数据
    输出格式：JSON字符串，包含以下字段：
    - success：操作是否成功
    - message：操作结果消息
    - data：文件内容（read操作）或文件列表（list操作）
    - file_path：操作的沙盒路径
    - timestamp：操作时间戳
    重要警告：任何文件操作**必须**通过本工具执行，**不允许**在没有工具调用的情况下声称执行了文件操作。
    """

    # 获取任务ID
    task_id = get_task_id()
    if not task_id:
        return json_module.dumps({
            "success": False,
            "message": "未找到任务ID，无法执行文件操作",
            "data": None,
            "file_path": "",
            "file_path_abs": "",
            "file_path_rel": "",
            "timestamp": datetime.now().isoformat()
        }, ensure_ascii=False)
    
    # 确保路径在沙盒中
    sandbox_file_path = _ensure_sandbox_path(task_id, file_path)
    sandbox_file_path_abs = str(sandbox_file_path)
    sandbox_file_path_rel = _sandbox_relative_path(task_id, sandbox_file_path)
    
    # 验证文件扩展名（除了list操作）
    if action != "list" and not _validate_file_extension(file_path):
        return json_module.dumps({
            "success": False,
            "message": f"不支持的文件格式：{Path(file_path).suffix}",
            "data": None,
            "file_path": sandbox_file_path_abs,
            "file_path_abs": sandbox_file_path_abs,
            "file_path_rel": sandbox_file_path_rel,
            "timestamp": datetime.now().isoformat()
        }, ensure_ascii=False)
    
    try:
        if action == "create":
            # 确保父目录存在
            sandbox_file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 写入文件内容
            ext = Path(file_path).suffix.lower()
            if ext == '.json' and data:
                with open(sandbox_file_path, 'w', encoding='utf-8', errors='replace') as f:
                    json_module.dump(data, f, ensure_ascii=False, indent=2)
            elif ext == '.docx':
                doc = Document()
                if content:
                    for line in content.split('\n'):
                        doc.add_paragraph(line)
                doc.save(sandbox_file_path)
            elif ext in {'.pdf', '.xlsx'} and content:
                try:
                    raw_bytes = base64.b64decode(content.encode("utf-8"), validate=False)
                except Exception as e:
                    return json_module.dumps({
                        "success": False,
                        "message": f"二进制文件内容需要 base64 编码字符串，解码失败：{str(e)}",
                        "data": None,
                        "file_path": sandbox_file_path_abs,
                        "file_path_abs": sandbox_file_path_abs,
                        "file_path_rel": sandbox_file_path_rel,
                        "timestamp": datetime.now().isoformat()
                    }, ensure_ascii=False)
                with open(sandbox_file_path, 'wb') as f:
                    f.write(raw_bytes)
            elif content:
                with open(sandbox_file_path, 'w', encoding='utf-8', errors='replace') as f:
                    f.write(content)
            else:
                return json_module.dumps({
                    "success": False,
                    "message": "创建文件需要提供content或data参数",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            return json_module.dumps({
                "success": True,
                "message": f"文件创建成功：{str(sandbox_file_path)}",
                "data": None,
                "file_path": sandbox_file_path_abs,
                "file_path_abs": sandbox_file_path_abs,
                "file_path_rel": sandbox_file_path_rel,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
        
        elif action == "read":
            if not sandbox_file_path.exists():
                return json_module.dumps({
                    "success": False,
                    "message": f"文件不存在：{str(sandbox_file_path)}",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            # 读取文件内容
            try:
                ext = Path(file_path).suffix.lower()
                if ext == '.json':
                    with open(sandbox_file_path, 'r', encoding='utf-8', errors='replace') as f:
                        file_content = json_module.load(f)
                elif ext == '.docx':
                    doc = Document(sandbox_file_path)
                    paragraphs = []
                    for para in doc.paragraphs:
                        paragraphs.append(para.text)
                    file_content = '\n'.join(paragraphs)
                elif ext == '.csv':
                    total_bytes = sandbox_file_path.stat().st_size
                    max_preview_bytes = 20000
                    max_preview_lines = 80
                    buf: List[str] = []
                    read_bytes = 0
                    truncated = False
                    with open(sandbox_file_path, 'r', encoding='utf-8', errors='replace') as f:
                        for _ in range(max_preview_lines):
                            line = f.readline()
                            if not line:
                                break
                            buf.append(line.rstrip("\n"))
                            read_bytes += len(line.encode("utf-8", errors="ignore"))
                            if read_bytes >= max_preview_bytes:
                                truncated = True
                                break
                        if not truncated:
                            if f.read(1):
                                truncated = True
                    file_content = "\n".join(buf)
                    return json_module.dumps({
                        "success": True,
                        "message": f"CSV 文件已返回预览内容：{str(sandbox_file_path)}",
                        "data": file_content,
                        "file_path": sandbox_file_path_abs,
                        "file_path_abs": sandbox_file_path_abs,
                        "file_path_rel": sandbox_file_path_rel,
                        "truncated": truncated or (total_bytes > read_bytes),
                        "total_bytes": int(total_bytes),
                        "preview_bytes": int(read_bytes),
                        "timestamp": datetime.now().isoformat()
                    }, ensure_ascii=False)
                elif ext in {'.pdf', '.xlsx'}:
                    with open(sandbox_file_path, 'rb') as f:
                        b = f.read()
                    file_content = base64.b64encode(b).decode("utf-8")
                else:
                    with open(sandbox_file_path, 'r', encoding='utf-8', errors='replace') as f:
                        file_content = f.read()
            except Exception as e:
                return json_module.dumps({
                    "success": False,
                    "message": f"读取文件失败：{str(e)}",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            return json_module.dumps({
                "success": True,
                "message": f"文件读取成功：{str(sandbox_file_path)}",
                "data": file_content,
                "file_path": sandbox_file_path_abs,
                "file_path_abs": sandbox_file_path_abs,
                "file_path_rel": sandbox_file_path_rel,
                "data_encoding": "base64" if Path(file_path).suffix.lower() in {".pdf", ".xlsx"} else "",
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
        
        elif action == "update":
            if not sandbox_file_path.exists():
                return json_module.dumps({
                    "success": False,
                    "message": f"文件不存在：{str(sandbox_file_path)}",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            # 更新文件内容
            ext = Path(file_path).suffix.lower()
            if ext == '.json' and data:
                with open(sandbox_file_path, 'w', encoding='utf-8', errors='replace') as f:
                    json_module.dump(data, f, ensure_ascii=False, indent=2)
            elif ext == '.docx':
                # 对于docx文件，我们创建一个新文档并替换原文件
                doc = Document()
                if content:
                    for line in content.split('\n'):
                        doc.add_paragraph(line)
                doc.save(sandbox_file_path)
            elif ext in {'.pdf', '.xlsx'} and content:
                try:
                    raw_bytes = base64.b64decode(content.encode("utf-8"), validate=False)
                except Exception as e:
                    return json_module.dumps({
                        "success": False,
                        "message": f"二进制文件内容需要 base64 编码字符串，解码失败：{str(e)}",
                        "data": None,
                        "file_path": sandbox_file_path_abs,
                        "file_path_abs": sandbox_file_path_abs,
                        "file_path_rel": sandbox_file_path_rel,
                        "timestamp": datetime.now().isoformat()
                    }, ensure_ascii=False)
                with open(sandbox_file_path, 'wb') as f:
                    f.write(raw_bytes)
            elif content:
                with open(sandbox_file_path, 'w', encoding='utf-8', errors='replace') as f:
                    f.write(content)
            else:
                return json_module.dumps({
                    "success": False,
                    "message": "更新文件需要提供content或data参数",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            return json_module.dumps({
                "success": True,
                "message": f"文件更新成功：{str(sandbox_file_path)}",
                "data": None,
                "file_path": sandbox_file_path_abs,
                "file_path_abs": sandbox_file_path_abs,
                "file_path_rel": sandbox_file_path_rel,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
        
        elif action == "delete":
            if not sandbox_file_path.exists():
                return json_module.dumps({
                    "success": False,
                    "message": f"文件不存在：{str(sandbox_file_path)}",
                    "data": None,
                    "file_path": sandbox_file_path_abs,
                    "file_path_abs": sandbox_file_path_abs,
                    "file_path_rel": sandbox_file_path_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            # 删除文件
            sandbox_file_path.unlink()
            
            return json_module.dumps({
                "success": True,
                "message": f"文件删除成功：{str(sandbox_file_path)}",
                "data": None,
                "file_path": sandbox_file_path_abs,
                "file_path_abs": sandbox_file_path_abs,
                "file_path_rel": sandbox_file_path_rel,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
        
        elif action == "list":
            # 确定要列出的目录
            target_path = sandbox_file_path
            if not target_path.exists():
                target_abs = str(target_path)
                target_rel = _sandbox_relative_path(task_id, target_path)
                return json_module.dumps({
                    "success": False,
                    "message": f"目录不存在：{str(target_path)}",
                    "data": None,
                    "file_path": target_abs,
                    "file_path_abs": target_abs,
                    "file_path_rel": target_rel,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)
            
            if not target_path.is_dir():
                target_path = target_path.parent
            
            # 列出目录中的文件
            files = []
            for item in target_path.iterdir():
                files.append({
                    "name": item.name,
                    "path": str(item),
                    "is_dir": item.is_dir(),
                    "size": item.stat().st_size if item.is_file() else 0,
                    "modified": datetime.fromtimestamp(item.stat().st_mtime).isoformat()
                })
            
            target_abs = str(target_path)
            target_rel = _sandbox_relative_path(task_id, target_path)
            return json_module.dumps({
                "success": True,
                "message": f"目录列出成功：{str(target_path)}",
                "data": files,
                "file_path": target_abs,
                "file_path_abs": target_abs,
                "file_path_rel": target_rel,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
        
        else:
            return json_module.dumps({
                "success": False,
                "message": f"不支持的操作类型：{action}",
                "data": None,
                "file_path": sandbox_file_path_abs,
                "file_path_abs": sandbox_file_path_abs,
                "file_path_rel": sandbox_file_path_rel,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)
    
    except Exception as e:
        return json_module.dumps({
            "success": False,
            "message": f"操作失败：{str(e)}",
            "data": None,
            "file_path": sandbox_file_path_abs if "sandbox_file_path_abs" in locals() else str(sandbox_file_path),
            "file_path_abs": sandbox_file_path_abs if "sandbox_file_path_abs" in locals() else str(sandbox_file_path),
            "file_path_rel": sandbox_file_path_rel if "sandbox_file_path_rel" in locals() else "",
            "timestamp": datetime.now().isoformat()
        }, ensure_ascii=False)
